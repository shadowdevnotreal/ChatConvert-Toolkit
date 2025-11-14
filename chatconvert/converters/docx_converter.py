"""
DOCX Converter - Create Microsoft Word documents from conversations.

Features:
- Professional document styling
- Color-coded participants
- Date grouping
- Statistics tables
- Page headers and footers
- Table of contents
- Formatted paragraphs with proper spacing
"""

import time
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .base_converter import BaseConverter
from ..models import Conversation, ConversionResult


class DOCXConverter(BaseConverter):
    """Convert conversations to Microsoft Word DOCX format."""

    def get_file_extension(self) -> str:
        return 'docx'

    def convert(self, conversation: Conversation, output_file: str) -> ConversionResult:
        """
        Convert conversation to DOCX format.

        Config options:
        - include_toc: Include table of contents (default: False)
        - include_stats: Include statistics table (default: True)
        - group_by_date: Group messages by date (default: True)
        - color_coded: Color-code participants (default: True)
        - page_break_dates: Add page break between dates (default: False)

        Args:
            conversation: Conversation to convert
            output_file: Path to output .docx file

        Returns:
            ConversionResult

        Raises:
            ImportError: If python-docx is not installed
        """
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX conversion. Install with: pip install python-docx")

        start_time = time.time()
        path = self._validate_output_path(output_file)

        try:
            # Get config
            include_toc = self.config.get('include_toc', False)
            include_stats = self.config.get('include_stats', True)
            group_by_date = self.config.get('group_by_date', True)
            color_coded = self.config.get('color_coded', True)
            page_break_dates = self.config.get('page_break_dates', False)

            # Generate DOCX
            self._generate_docx(
                conversation,
                str(path),
                include_toc=include_toc,
                include_stats=include_stats,
                group_by_date=group_by_date,
                color_coded=color_coded,
                page_break_dates=page_break_dates
            )

            processing_time = time.time() - start_time

            return self._create_result(
                success=True,
                output_file=str(path),
                message_count=len(conversation.messages),
                processing_time=processing_time,
                statistics={
                    'participants': len(conversation.get_participants_list()),
                    'include_toc': include_toc,
                }
            )

        except Exception as e:
            self.logger.error(f"DOCX conversion failed: {e}")
            return self._create_result(
                success=False,
                errors=[str(e)],
                processing_time=time.time() - start_time
            )

    def _generate_docx(
        self,
        conversation: Conversation,
        output_path: str,
        include_toc: bool = False,
        include_stats: bool = True,
        group_by_date: bool = True,
        color_coded: bool = True,
        page_break_dates: bool = False
    ):
        """Generate DOCX document."""

        # Create document
        doc = Document()

        # Set document properties
        doc.core_properties.title = conversation.title
        doc.core_properties.subject = f"{conversation.platform} Conversation"
        doc.core_properties.creator = "ChatConvert-Toolkit"

        # Color palette for participants
        participant_colors = [
            (102, 126, 234),  # #667eea
            (118, 75, 162),   # #764ba2
            (240, 147, 251),  # #f093fb
            (79, 172, 254),   # #4facfe
            (67, 233, 123),   # #43e97b
        ]

        # Map participants to colors
        participants = conversation.get_participants_list()
        participant_color_map = {}
        for i, p in enumerate(participants):
            participant_color_map[p] = participant_colors[i % len(participant_colors)]

        # Title
        title = doc.add_heading(f"💬 {conversation.title}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with metadata
        date_range = conversation.get_date_range()
        if date_range[0]:
            start = date_range[0].strftime('%Y-%m-%d')
            end = date_range[1].strftime('%Y-%m-%d')
            if start == end:
                date_str = f"Date: {start}"
            else:
                date_str = f"Period: {start} to {end}"
        else:
            date_str = ""

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{conversation.platform} Conversation\n{date_str}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacing

        # Table of contents
        if include_toc:
            doc.add_heading("Table of Contents", 1)
            toc = doc.add_paragraph()
            toc.add_run("[Table of Contents will be generated when opened in Microsoft Word]")
            toc.italic = True
            doc.add_page_break()

        # Statistics
        if include_stats:
            self._add_statistics_section(doc, conversation, participant_color_map)
            doc.add_page_break()

        # Messages
        if group_by_date:
            self._add_messages_by_date(
                doc, conversation, participant_color_map,
                color_coded, page_break_dates
            )
        else:
            self._add_messages_chronological(
                doc, conversation, participant_color_map, color_coded
            )

        # Footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by ChatConvert-Toolkit on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save document
        doc.save(output_path)

    def _add_statistics_section(self, doc, conversation: Conversation, color_map: dict):
        """Add statistics section to document."""

        doc.add_heading("📊 Conversation Statistics", 1)

        # Calculate stats
        participants = conversation.get_participants_list()
        participant_stats = []

        for p in participants:
            count = sum(1 for m in conversation.messages if m.sender == p)
            percentage = (count / len(conversation.messages) * 100) if conversation.messages else 0
            participant_stats.append((p, count, percentage))

        # Sort by message count
        participant_stats.sort(key=lambda x: x[1], reverse=True)

        # Create table
        table = doc.add_table(rows=len(participant_stats) + 1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Participant'
        header_cells[1].text = 'Messages'
        header_cells[2].text = 'Percentage'

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for i, (username, count, percentage) in enumerate(participant_stats, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = username
            row_cells[1].text = str(count)
            row_cells[2].text = f"{percentage:.1f}%"

            # Color-code participant name
            if username in color_map:
                r, g, b = color_map[username]
                row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g, b)
                row_cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()  # Spacing

        # Overall stats
        p = doc.add_paragraph()
        p.add_run(f"Total Messages: ").bold = True
        p.add_run(f"{len(conversation.messages)}\n")
        p.add_run(f"Total Participants: ").bold = True
        p.add_run(f"{len(participants)}")

    def _add_messages_by_date(
        self, doc, conversation: Conversation,
        color_map: dict, color_coded: bool, page_break_dates: bool
    ):
        """Add messages grouped by date."""

        current_date = None

        for msg in conversation.messages:
            date_str = msg.timestamp.strftime('%Y-%m-%d')

            # Add date header if changed
            if date_str != current_date:
                if current_date is not None and page_break_dates:
                    doc.add_page_break()

                current_date = date_str
                doc.add_heading(f"📅 {date_str}", 2)

            # Add message
            self._add_message(doc, msg, color_map, color_coded, show_date=False)

    def _add_messages_chronological(
        self, doc, conversation: Conversation,
        color_map: dict, color_coded: bool
    ):
        """Add messages in chronological order."""

        doc.add_heading("Messages", 1)

        for msg in conversation.messages:
            self._add_message(doc, msg, color_map, color_coded, show_date=True)

    def _add_message(
        self, doc, msg, color_map: dict,
        color_coded: bool, show_date: bool = True
    ):
        """Add a single message to the document."""

        # Message header (sender + timestamp)
        if show_date:
            timestamp = msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        else:
            timestamp = msg.timestamp.strftime('%H:%M:%S')

        p = doc.add_paragraph()
        run = p.add_run(f"{msg.sender}")
        run.font.bold = True
        run.font.size = Pt(11)

        # Color-code sender
        if color_coded and msg.sender in color_map:
            r, g, b = color_map[msg.sender]
            run.font.color.rgb = RGBColor(r, g, b)

        # Add timestamp
        time_run = p.add_run(f" - {timestamp}")
        time_run.font.size = Pt(9)
        time_run.font.color.rgb = RGBColor(128, 128, 128)

        # Message content
        content_p = doc.add_paragraph(msg.content)
        content_p.paragraph_format.left_indent = Inches(0.25)
        content_p.paragraph_format.space_after = Pt(6)
